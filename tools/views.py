import os
import mimetypes
import json
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.conf import settings
from docx import Document
from pathlib import Path 
from tools.serialized import FileSerializer, WriteFileSerializer

# Uso Path para manejar las rutas de forma más moderna y segura
FILE_BASE_PATH = Path(settings.BASE_DIR) / 'files'
# Crea el directorio si no existe
if not FILE_BASE_PATH.exists():
    FILE_BASE_PATH.mkdir()

# Define las extensiones de archivo permitidas
ALLOWED_EXTENSIONS = {'.txt', '.md', '.json', '.csv', '.docx'} 

def is_valid_file_path(filename: str) -> Path:
    """
    Valida que el archivo solicitado sea seguro y devuelva su ruta completa.
    Lanza una excepción si la validación falla.
    
    Esta función centraliza la lógica de seguridad y previene ataques de 
    recorrido de directorios (path traversal).
    """
    if not filename:
        raise ValueError("El nombre del archivo no puede estar vacío.")
    
    file_path = FILE_BASE_PATH / filename
    
    # 1. Previene ataques de path traversal: comprueba que la ruta resuelta 
    #    está dentro del directorio base.
    if not file_path.resolve().is_relative_to(FILE_BASE_PATH.resolve()):
        raise PermissionError("Acceso denegado. La ruta no es válida.")
        
    # 2. Valida la extensión del archivo para asegurar que es un tipo permitido.
    if file_path.suffix.lower() not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Extensión de archivo no permitida: {file_path.suffix}")

    return file_path


def get_tools_manifest():
    return {
        "read_file": {
            "description": "Reads the content of a specified text file.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "The name of the file to read, e.g., 'test.txt'."
                    }
                },
                "required": ["filename"]
            }
        },
        "write_file": {
            "description": "Writes content to a specified text file.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "The name of the file to write, e.g., 'new_file.txt'."
                    },
                    "content": {
                        "type": "string",
                        "description": "The content to write into the file."
                    }
                },
                "required": ["filename", "content"]
            }
        },
        "list_files": {
            "description": "Lists all the files in the 'files' directory.",
            "input_schema": {
                "type": "object",
                "properties": {}
            }
        },
        "read_docx_file": {
            "description": "Reads the content of a specified .docx file.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "The name of the .docx file to read, e.g., 'report.docx'."
                    }
                },
                "required": ["filename"]
            }
        }
    }



class ToolsManifestView(APIView):
    def get(self, request):
        return Response(get_tools_manifest())

class WriteFileView(APIView):
    def post(self, request):
        filename = request.data.get('filename')
        content = request.data.get('content')

        try:
            # Usa la función de validación y obtiene la ruta
            file_path = is_valid_file_path(filename)
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return Response({"message": f"Successfully wrote to file: {filename}"})
        except ValueError as e:
            # Error de validación del nombre o extensión
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)
        except PermissionError as e:
            # Error de seguridad
            return Response({"error": str(e)}, status=status.HTTP_403_FORBIDDEN)
        except Exception as e:
            # Otros errores inesperados
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class ReadFileView(APIView):
    def post(self, request):
        filename = request.data.get('filename')
        
        try:
            # Usa la función de validación y obtiene la ruta
            file_path = is_valid_file_path(filename)
            
            if file_path.suffix.lower() == '.docx':
                doc = Document(file_path)
                content = '\n'.join([p.text for p in doc.paragraphs])
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()

            return Response({"filename": filename, "content": content})
        except ValueError as e:
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)
        except PermissionError as e:
            return Response({"error": str(e)}, status=status.HTTP_403_FORBIDDEN)
        except FileNotFoundError:
            return Response({"error": "File not found."}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class ListFilesView(APIView):
    def get(self, request):
        try:
            # Lista los archivos y devuelve solo el nombre.
            files = [f.name for f in FILE_BASE_PATH.iterdir() if f.is_file() and f.suffix.lower() in ALLOWED_EXTENSIONS]
            return Response({"files": files})
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class WriteFileView(APIView):
    def post(self, request):
        serializer = WriteFileSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)

        filename = serializer.validated_data['filename']
        content = serializer.validated_data['content']

        try:
            file_path = is_valid_file_path(filename)

            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return Response({"message": f"Successfully wrote to file: {filename}"})
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class ReadFileView(APIView):
    def post(self, request):
        serializer = FileSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)

        filename = serializer.validated_data['filename']

        try:
            file_path = is_valid_file_path(filename)

            if file_path.suffix.lower() == '.docx':
                doc = Document(file_path)
                content = '\n'.join([p.text for p in doc.paragraphs])
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()

            return Response({"filename": filename, "content": content})
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
